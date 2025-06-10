# extract_text.py
import os
import pdfplumber
import pypdfium2 as pdfium
from pdfminer.high_level import extract_text as pdfminer_extract
from docx import Document
import gc
import json

def extract_with_pdfplumber(path):
    """Thử extract bằng pdfplumber"""
    try:
        full_text = ""
        with pdfplumber.open(path) as pdf:
            total_pages = len(pdf.pages)
            print(f"  → [pdfplumber] PDF có {total_pages} trang")

            for i, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + "\n"
                        print(f"  → Trang {i}/{total_pages}: Extract thành công")
                    else:
                        print(f"  → Trang {i}/{total_pages}: Không extract được text")
                except Exception as e:
                    print(f"  → Lỗi ở trang {i}/{total_pages}: {str(e)}")

        return full_text.strip()
    except Exception as e:
        print(f"  → Lỗi pdfplumber: {str(e)}")
        return ""


def extract_with_pdfium(path):
    """Thử extract bằng pypdfium2"""
    try:
        pdf = pdfium.PdfDocument(path)
        full_text = ""
        total_pages = len(pdf)
        print(f"  → [pdfium] PDF có {total_pages} trang")

        for i, page in enumerate(pdf, 1):
            try:
                textpage = page.get_textpage()
                text = textpage.get_text_range()
                if text:
                    full_text += text + "\n"
                    print(f"  → Trang {i}/{total_pages}: Extract thành công")
                else:
                    print(f"  → Trang {i}/{total_pages}: Không extract được text")
            except Exception as e:
                print(f"  → Lỗi ở trang {i}/{total_pages}: {str(e)}")

        return full_text.strip()
    except Exception as e:
        print(f"  → Lỗi pdfium: {str(e)}")
        return ""


def extract_with_pdfminer(path):
    """Thử extract bằng pdfminer"""
    try:
        print("  → [pdfminer] Đang extract...")
        text = pdfminer_extract(path)
        if text.strip():
            print("  → Extract thành công")
            return text.strip()
        else:
            print("  → Không extract được text")
            return ""
    except Exception as e:
        print(f"  → Lỗi pdfminer: {str(e)}")
        return ""


def extract_text_from_docx(path):
    """Extract text from a .docx file."""
    try:
        doc = Document(path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        return full_text.strip()
    except Exception as e:
        print(f"  → Lỗi extract docx: {str(e)}")
        return ""


def extract_text_from_pdf(path):
    """
    Thử extract text bằng nhiều phương pháp khác nhau cho đến khi thành công
    """
    results = []

    # 1. Thử pdfplumber trước
    print("\n→ Thử phương pháp 1: pdfplumber")
    text = extract_with_pdfplumber(path)
    if text:
        results.append(("pdfplumber", text))

    # 2. Thử pdfium
    print("\n→ Thử phương pháp 2: pypdfium2")
    text = extract_with_pdfium(path)
    if text:
        results.append(("pdfium", text))

    # 3. Thử pdfminer với các encoding khác nhau
    print("\n→ Thử phương pháp 3: pdfminer")
    encodings = ['utf-8', 'utf-16', 'windows-1258', 'latin1']
    for encoding in encodings:
        try:
            text = pdfminer_extract(path, codec=encoding)
            if text.strip():
                print(f"  → Extract thành công với encoding {encoding}")
                results.append(("pdfminer", text))
                break
        except Exception as e:
            print(f"  → Lỗi với encoding {encoding}: {str(e)}")
            continue

    # Chọn kết quả tốt nhất
    if results:
        # Sắp xếp theo độ dài text (ưu tiên text dài hơn)
        results.sort(key=lambda x: len(x[1]), reverse=True)
        print(f"\n→ Thành công! Phương pháp tốt nhất: {results[0][0]}")
        return results[0][1]

    print("\n→ Tất cả phương pháp đều thất bại!")
    return ""


def extract_text_from_file(path):
    """
    Extract text từ file PDF, DOCX.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        print("\n→ Thử extract file Word (.docx)")
        return extract_text_from_docx(path)
    else:
        print(f"→ Không hỗ trợ định dạng file: {ext}")
        return ""

def detect_field(folder):
    folder = folder.lower()
    # Nhận diện lĩnh vực theo tên folder gốc
    if "criminal" in folder:
        return "Hình sự"
    if "civil" in folder:
        return "Dân sự"
    if "administrative" in folder:
        return "Hành chính"
    return ""


def is_law_text(text: str) -> bool:
    """
    Kiểm tra xem text có phải là văn bản luật thực sự không.
    Trả về True nếu có các từ khóa đặc trưng của luật/pháp lý.
    """
    # Các từ khóa đặc trưng của văn bản luật
    keywords = [
        "điều", "khoản", "luật", "nghị định", "quyết định", "thông tư",
        "bộ luật", "chương", "mục", "căn cứ", "quốc hội", "chính phủ", "ủy ban", "bộ trưởng"
    ]
    text_lower = text.lower()
    # Nếu có ít nhất 2 từ khóa xuất hiện thì coi là hợp lệ
    count = sum(1 for kw in keywords if kw in text_lower)
    return count >= 2


if __name__ == "__main__":
    # Danh sách các folder cần xử lý
    folders = ["law documents/civil", "law documents/criminal"]
    output_folder = "./plain_texts"
    os.makedirs(output_folder, exist_ok=True)

    all_failed_files = []
    for folder in folders:
        if not os.path.isdir(folder):
            print(f"⚠️ Folder không tồn tại: {folder}")
            continue

        # Lấy danh sách file PDF, DOCX, DOC chưa được extract
        failed_files = [
            f for f in os.listdir(folder)
            if (
                (f.lower().endswith(".pdf") or f.lower().endswith(".docx"))
                and not os.path.exists(os.path.join(output_folder, f.replace(".pdf", ".txt").replace(".docx", ".txt")))
            )
        ]

        print(f"Tìm thấy {len(failed_files)} file PDF/DOCX chưa được extract trong {folder}:")
        for f in failed_files:
            print(f"  - {f}")

        all_failed_files.extend([(folder, f) for f in failed_files])

    if not all_failed_files:
        print("Tất cả file đã được extract!")
        exit(0)

    # Chuẩn bị dict lưu mapping {filename gốc: field}
    fields_meta_path = os.path.join(output_folder, "fields.meta")
    if os.path.exists(fields_meta_path):
        with open(fields_meta_path, "r", encoding="utf-8") as f:
            fields_dict = json.load(f)
    else:
        fields_dict = {}

    success = 0
    failed = 0

    for folder, filename in all_failed_files:
        print(f"\nXử lý: {filename} (folder: {folder})")
        file_path = os.path.abspath(os.path.join(folder, filename))
        text = extract_text_from_file(file_path)

        # Lọc file không chứa văn bản luật thực sự
        if not text or not is_law_text(text):
            print("  ❌ Bỏ qua: không có text hợp lệ để lưu (không phải văn bản luật)")
            gc.collect()
            # Xóa file gốc trong law documents
            try:
                os.remove(file_path)
                print(f"  ⚠️ Đã xóa file không hợp lệ: {file_path}")
            except Exception as e:
                print(f"  ⚠️ Lỗi khi xóa file: {str(e)}")
            # Xóa file .txt đã extract (nếu có) trong plain_texts
            out_path = os.path.join(output_folder, filename.replace(".pdf", ".txt").replace(".docx", ".txt"))
            if os.path.exists(out_path):
                try:
                    os.remove(out_path)
                    print(f"  ⚠️ Đã xóa file extract: {out_path}")
                except Exception as e:
                    print(f"  ⚠️ Lỗi khi xóa file extract: {str(e)}")
            failed += 1
            continue

        out_path = os.path.join(output_folder, filename.replace(".pdf", ".txt").replace(".docx", ".txt"))
        field = detect_field(folder)
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
            # Cập nhật dict fields nếu xác định được lĩnh vực
            if field:
                fields_dict[os.path.basename(out_path)] = field
                print(f"  ✓ Đã cập nhật metadata lĩnh vực cho {os.path.basename(out_path)}: {field}")
            else:
                print("  ⚠️ Không xác định được lĩnh vực, không cập nhật fields.meta")
            print(f"  ✓ Đã lưu: {out_path}")
            success += 1
        except Exception as e:
            print(f"  ❌ Lỗi lưu file: {str(e)}")
            failed += 1

    # Ghi lại file fields.meta duy nhất (dạng JSON)
    with open(fields_meta_path, "w", encoding="utf-8") as f:
        json.dump(fields_dict, f, ensure_ascii=False, indent=2)

    print(f"\n=== Tổng kết ===")
    print(f"Thành công: {success} file")
    print(f"Thất bại: {failed} file")
